import os
import json
import uuid
import pytest
from fastapi.testclient import TestClient
from app.main import app
from app.core.config import get_config
from app.core.database import init_db

client = TestClient(app)


@pytest.fixture(scope="module")
def setup_test_env():
    # 确保存放测试文件的目录存在
    cfg = get_config()
    os.makedirs(cfg.get("files_dir", "data/files"), exist_ok=True)
    # 初始化数据库
    init_db(cfg)

    # 模拟登录获取 token
    login_resp = client.post(
        "/api/auth/login", json={"username": "admin", "password": "adminpassword"})
    if login_resp.status_code != 200:
        # 如果admin不存在，先注册一个
        client.post("/api/auth/register", json={"username": "admin",
                    "password": "adminpassword", "email": "admin@example.com"})
        login_resp = client.post(
            "/api/auth/login", json={"username": "admin", "password": "adminpassword"})

    token = login_resp.json().get("access_token")
    headers = {"Authorization": f"Bearer {token}"}

    from docx import Document

    # 创建测试用的法规文件和合同文件
    reg_path = "tests/test_tax_reg.docx"
    doc = Document()
    doc.add_paragraph(
        "第一条 增值税税率为13%。\n第二条 企业所得税税率为25%。\n第三条 纳税人应当在收到发票后30日内付款。")
    doc.save(reg_path)

    contract_path = "tests/test_tax_contract.docx"
    doc2 = Document()
    doc2.add_paragraph(
        "第一章 税务条款\n1.1 本合同约定的增值税税率为9%。\n1.2 甲方应在收到发票后45日内付款。\n1.3 乙方需提供增值税专用发票。")
    doc2.save(contract_path)

    yield headers, reg_path, contract_path

    # 清理测试文件
    if os.path.exists(reg_path):
        os.remove(reg_path)
    if os.path.exists(contract_path):
        os.remove(contract_path)


def test_tax_audit_full_pipeline(setup_test_env):
    headers, reg_path, contract_path = setup_test_env

    # 1. 上传并解析法规
    with open(reg_path, "rb") as f:
        resp = client.post("/tax-audit/regulations/import",
                           files={"file": ("reg.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}, headers=headers)
    assert resp.status_code == 200
    reg_doc_id = resp.json()["document_id"]

    resp = client.post(
        f"/tax-audit/regulations/{reg_doc_id}/parse", headers=headers)
    assert resp.status_code == 200
    assert resp.json()["rule_count"] > 0
    print(f"\n[1] 法规解析完成，提取规则数: {resp.json()['rule_count']}")

    # 2. 上传并解析合同
    with open(contract_path, "rb") as f:
        resp = client.post("/tax-audit/contracts/import",
                           files={"file": ("contract.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}, headers=headers)
    assert resp.status_code == 200
    contract_id = resp.json()["contract_id"]

    resp = client.post(
        f"/tax-audit/contracts/{contract_id}/analyze", headers=headers)
    assert resp.status_code == 200
    assert resp.json()["clause_count"] > 0
    print(f"[2] 合同解析完成，切分条款数: {resp.json()['clause_count']}")

    # 3. 执行匹配 (Match)
    resp = client.post(
        f"/tax-audit/contracts/{contract_id}/match", headers=headers)
    assert resp.status_code == 200
    print(
        f"[3] 匹配执行完成，共生成 {resp.json()['total_matches']} 条匹配记录 (合规: {resp.json()['compliant_count']}, 违规: {resp.json()['non_compliant_count']})")

    # 4. 生成审计风险项 (Generate Issues)
    resp = client.post(
        f"/tax-audit/contracts/{contract_id}/issues/generate", headers=headers)
    assert resp.status_code == 200
    print(
        f"[4] 风险项生成完成，共 {resp.json()['total']} 项 (高风险: {resp.json()['high']})")

    # 5. 获取风险项列表并进行人工复核 (Review)
    resp = client.get(
        f"/tax-audit/contracts/{contract_id}/issues", headers=headers)
    assert resp.status_code == 200
    issues = resp.json()["items"]
    assert len(issues) > 0

    issue_to_review = issues[0]
    issue_id = issue_to_review["id"]
    print(
        f"[5] 准备复核风险项 ID: {issue_id}, 原状态: {issue_to_review['reviewer_status']}, 风险描述: {issue_to_review['issue_text']}")

    review_payload = {
        "reviewer_status": "rejected",
        "reviewer_note": "业务特批，允许税率为9%",
        "risk_level": "low"
    }
    resp = client.post(
        f"/tax-audit/issues/{issue_id}/review", json=review_payload, headers=headers)
    assert resp.status_code == 200
    assert resp.json()["reviewer_status"] == "rejected"
    print(
        f"[6] 人工复核完成，新状态: {resp.json()['reviewer_status']}, 备注: {resp.json()['reviewer_note']}")

    # 6. 验证审计轨迹 (Audit Trace)
    resp = client.get(f"/tax-audit/issues/{issue_id}/trace", headers=headers)
    assert resp.status_code == 200
    traces = resp.json()["items"]
    assert len(traces) > 0
    print(
        f"[7] 查找到该风险项的审计轨迹共 {len(traces)} 条。最新动作: {traces[0]['action_type']}")
