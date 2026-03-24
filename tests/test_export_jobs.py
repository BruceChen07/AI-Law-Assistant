import os
import uuid
from datetime import datetime, timezone
from docx import Document
from docx.oxml.ns import qn
from app.core.database import init_db, get_conn
from app.services.crud import (
    create_tax_regulation_document,
    create_tax_contract_document,
    replace_contract_clauses,
)
from app.services.tax_matcher import match_contract_against_rules
from app.services.tax_risk import generate_issues_from_matches
from app.services.tax_report import export_tax_audit_report
from app.services.export_jobs import submit_tax_report_export_job, get_tax_report_export_job


def _seed_data(cfg, tmp_path):
    reg_id = str(uuid.uuid4())
    contract_id = str(uuid.uuid4())
    create_tax_regulation_document(
        cfg=cfg,
        document_id=reg_id,
        original_filename="reg.pdf",
        file_path=str(tmp_path / "reg.pdf"),
        file_type="pdf",
        file_size=100,
        uploaded_by="u1",
        checksum="x1",
        parse_status="done",
    )
    create_tax_contract_document(
        cfg=cfg,
        document_id=contract_id,
        original_filename="contract.docx",
        file_path=str(tmp_path / "contract.docx"),
        file_type="docx",
        file_size=200,
        uploaded_by="u1",
        parse_status="done",
        ocr_used=0,
    )
    replace_contract_clauses(
        cfg,
        contract_id,
        [
            {"clause_path": "1.1", "page_no": 1, "paragraph_no": "1",
                "clause_text": "税率按9%执行", "entities_json": "{}"},
            {"clause_path": "1.2", "page_no": 2, "paragraph_no": "2",
                "clause_text": "双方按法规办理", "entities_json": "{}"},
        ],
        created_by="u1",
    )
    conn = get_conn(cfg)
    cur = conn.cursor()
    now = datetime.now(timezone.utc).isoformat()
    cur.execute(
        """
        INSERT INTO tax_rule(
            id, regulation_document_id, law_title, article_no, rule_type,
            trigger_condition, required_action, prohibited_action, numeric_constraints,
            deadline_constraints, region, industry, effective_date, expiry_date,
            source_page, source_paragraph, source_text, created_by, created_at, updated_at
        )
        VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
        """,
        (
            str(uuid.uuid4()),
            reg_id,
            "示例法规",
            "第一条",
            "tax_rate",
            "",
            "",
            "",
            "13%",
            "",
            "",
            "",
            "",
            "",
            1,
            "1",
            "增值税税率13%",
            "u1",
            now,
            now,
        ),
    )
    conn.commit()
    conn.close()
    match_contract_against_rules(
        cfg, contract_id, operator_id="u1", top_k_per_clause=1)
    generate_issues_from_matches(cfg, contract_id, operator_id="u1")
    return contract_id


def test_export_job_docx_and_idempotency(tmp_path):
    db_path = tmp_path / "test.db"
    files_dir = tmp_path / "files"
    files_dir.mkdir(parents=True, exist_ok=True)
    cfg = {"db_path": str(db_path), "files_dir": str(files_dir)}
    init_db(cfg)
    contract_id = _seed_data(cfg, tmp_path)

    job1 = submit_tax_report_export_job(
        cfg=cfg,
        contract_id=contract_id,
        export_format="docx",
        template_version="v1.0",
        locale="zh-CN",
        include_appendix=True,
        brand="",
        requester="u1",
    )
    assert job1["status"] == "done"
    assert os.path.exists(job1["output_path"])
    assert job1["output_path"].endswith(".docx")

    job2 = submit_tax_report_export_job(
        cfg=cfg,
        contract_id=contract_id,
        export_format="docx",
        template_version="v1.0",
        locale="zh-CN",
        include_appendix=True,
        brand="",
        requester="u1",
    )
    assert job2["export_id"] == job1["export_id"]

    fetched = get_tax_report_export_job(cfg, job1["export_id"])
    assert fetched is not None
    assert fetched["status"] == "done"

    conn = get_conn(cfg)
    cur = conn.cursor()
    cur.execute("SELECT COUNT(1) AS c FROM export_snapshot")
    snapshot_count = int(cur.fetchone()["c"])
    cur.execute("SELECT COUNT(1) AS c FROM evidence_anchor")
    anchor_count = int(cur.fetchone()["c"])
    conn.close()
    assert snapshot_count >= 1
    assert anchor_count >= 1


def test_docx_renderer_layout_and_content(tmp_path):
    db_path = tmp_path / "test.db"
    files_dir = tmp_path / "files"
    files_dir.mkdir(parents=True, exist_ok=True)
    cfg = {"db_path": str(db_path), "files_dir": str(files_dir)}
    init_db(cfg)
    contract_id = _seed_data(cfg, tmp_path)

    exported = export_tax_audit_report(
        cfg, contract_id, export_format="docx", template_version="v1.2", locale="zh-CN")
    assert exported["export_format"] == "docx"
    assert os.path.exists(exported["file_path"])

    doc = Document(exported["file_path"])
    text = "\n".join([p.text for p in doc.paragraphs if p.text])
    assert "合同审计报告" in text
    assert "目录" in text
    assert "一、总体概览" in text
    assert "三、风险清单" in text
    assert len(doc.tables) >= 1
    first_table = doc.tables[0]
    first_row = first_table.rows[0]._tr
    tr_pr = first_row.trPr
    assert tr_pr is not None
    assert tr_pr.find(qn("w:tblHeader")) is not None
    for cell in first_table.rows[0].cells:
        tc_pr = cell._tc.tcPr
        assert tc_pr is not None
        shd = tc_pr.find(qn("w:shd"))
        assert shd is not None
        assert shd.get(qn("w:fill")) == "D9E1F2"
    header_text = doc.sections[0].header.paragraphs[0].text
    footer_text = doc.sections[0].footer.paragraphs[0].text
    assert "合同ID" in header_text
    assert "模板版本: v1.2" in footer_text
    xml_text = doc.element.xml
    assert "fldSimple" in xml_text
    assert "TOC" in xml_text


def test_docx_renderer_locale_en(tmp_path):
    db_path = tmp_path / "test.db"
    files_dir = tmp_path / "files"
    files_dir.mkdir(parents=True, exist_ok=True)
    cfg = {"db_path": str(db_path), "files_dir": str(files_dir)}
    init_db(cfg)
    contract_id = _seed_data(cfg, tmp_path)

    exported = export_tax_audit_report(
        cfg, contract_id, export_format="docx", template_version="v2.0", locale="en-US", brand="ACME")
    assert exported["export_format"] == "docx"
    assert os.path.exists(exported["file_path"])

    doc = Document(exported["file_path"])
    text = "\n".join([p.text for p in doc.paragraphs if p.text])
    assert "Contract Audit Report" in text
    assert "Table of Contents" in text
    assert "I. Overview" in text
    assert "V. Risk Details" in text
    header_text = doc.sections[0].header.paragraphs[0].text
    footer_text = doc.sections[0].footer.paragraphs[0].text
    assert "Contract ID" in header_text
    assert "ACME" in header_text
    assert "Template Version: v2.0" in footer_text
