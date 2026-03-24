import os
import zipfile
from docx import Document
from app.services.docx_modifier import prepare_docx_for_comments, insert_risk_comments


def test_prepare_docx_for_comments_injects_required_files(tmp_path):
    # 1. 构造一个不含批注的原始 docx
    input_path = str(tmp_path / "original.docx")
    doc = Document()
    doc.add_paragraph("这是测试合同正文第一段。")
    doc.save(input_path)

    # 2. 运行 M1 骨架
    output_path = str(tmp_path / "output.docx")
    prepare_docx_for_comments(input_path, output_path)

    assert os.path.exists(output_path)

    # 3. 验证内部文件结构
    with zipfile.ZipFile(output_path, 'r') as z:
        file_list = z.namelist()

        # 必须注入 comments.xml
        assert 'word/comments.xml' in file_list
        comments_content = z.read('word/comments.xml').decode('utf-8')
        assert 'w:comments' in comments_content

        # 必须更新 [Content_Types].xml
        content_types = z.read('[Content_Types].xml').decode('utf-8')
        assert 'word/comments.xml' in content_types

        # 必须更新 word/_rels/document.xml.rels
        rels = z.read('word/_rels/document.xml.rels').decode('utf-8')
        assert 'comments.xml' in rels

    # 4. 确保 Word 仍然能正常读取它（无损证明）
    doc_out = Document(output_path)
    assert len(doc_out.paragraphs) == 1
    assert doc_out.paragraphs[0].text == "这是测试合同正文第一段。"


def test_insert_risk_comments_injects_comments_into_docx(tmp_path):
    # 1. 构造原始 docx
    input_path = str(tmp_path / "original_m2.docx")
    doc = Document()
    doc.add_paragraph("这是测试合同正文第一段，包含了需要批注的违约责任。")
    doc.add_paragraph("这是第二段正常内容。")
    doc.save(input_path)

    # 2. 准备风险数据
    risks = [
        {
            "issue_id": "r1",
            "risk_level": "high",
            "issue_text": "违约金比例过高",
            "suggestion": "建议降低到总价的20%以内",
            "evidence_text": "违约责任"
        }
    ]

    # 3. 运行 M2 批注注入
    output_path = str(tmp_path / "output_m2.docx")
    insert_risk_comments(input_path, output_path, risks)

    assert os.path.exists(output_path)

    # 4. 验证内部文件结构
    with zipfile.ZipFile(output_path, 'r') as z:
        comments_content = z.read('word/comments.xml').decode('utf-8')
        assert '违约金比例过高' in comments_content
        assert '建议降低到总价的20%以内' in comments_content
        assert 'w:comment ' in comments_content

        document_content = z.read('word/document.xml').decode('utf-8')
        assert 'w:commentRangeStart' in document_content
        assert 'w:commentRangeEnd' in document_content
        assert 'w:commentReference' in document_content

    # 5. 确保 Word 能正常读取修改后的文档
    doc_out = Document(output_path)
    assert len(doc_out.paragraphs) == 2
    assert "这是测试合同正文第一段" in doc_out.paragraphs[0].text
