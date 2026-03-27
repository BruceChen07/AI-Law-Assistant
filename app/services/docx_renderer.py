from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.oxml.ns import nsdecls
from docx.shared import Cm, Pt, RGBColor


def _labels(locale: str):
    if str(locale or "").lower().startswith("en"):
        return {
            "report_title": "Contract Audit Report",
            "toc": "Table of Contents",
            "toc_hint": "Right click this area in Word and choose Update Field.",
            "overview": "I. Overview",
            "summary": "II. Risk Summary",
            "risks": "III. Risk List",
            "appendix": "IV. Evidence Appendix",
            "header_title": "Contract Audit Export",
            "contract_id": "Contract ID",
            "generated_at": "Generated At",
            "template_version": "Template Version",
            "contract_file": "Contract File",
            "parse_status": "Parse Status",
            "clause_count": "Clause Count",
            "issue_count": "Risk Count",
            "trace_count": "Trace Count",
            "high": "High",
            "medium": "Medium",
            "low": "Low",
            "pending": "Pending Review",
            "confirmed": "Confirmed",
            "table_no": "No.",
            "table_level": "Risk Level",
            "table_issue": "Risk Description",
            "table_evidence": "Evidence Text",
            "table_suggestion": "Suggestion",
            "risk_details": "V. Risk Details",
            "risk_no": "Risk",
            "risk_level": "Risk Level",
            "risk_text": "Issue",
            "risk_suggestion": "Suggestion",
            "evidence_title": "Evidence",
            "evidence_issue_id": "Risk ID",
            "evidence_law": "Law",
            "evidence_loc": "Location",
            "evidence_text": "Source",
        }
    return {
        "report_title": "合同审计报告",
        "toc": "目录",
        "toc_hint": "请在 Word 中右键此区域并更新目录（引用 -> 目录）。",
        "overview": "一、总体概览",
        "summary": "二、风险统计",
        "risks": "三、风险清单",
        "appendix": "四、证据附录",
        "header_title": "合同审核导出报告",
        "contract_id": "合同ID",
        "generated_at": "生成时间",
        "template_version": "模板版本",
        "contract_file": "合同文件",
        "parse_status": "解析状态",
        "clause_count": "条款总数",
        "issue_count": "风险总数",
        "trace_count": "审计轨迹",
        "high": "高风险",
        "medium": "中风险",
        "low": "低风险",
        "pending": "待复核",
        "confirmed": "已确认",
        "table_no": "序号",
        "table_level": "风险等级",
        "table_issue": "风险描述",
        "table_evidence": "证据原文",
        "table_suggestion": "整改建议",
        "risk_details": "五、风险详情",
        "risk_no": "风险项",
        "risk_level": "风险等级",
        "risk_text": "问题描述",
        "risk_suggestion": "整改建议",
        "evidence_title": "证据",
        "evidence_issue_id": "关联风险ID",
        "evidence_law": "法规",
        "evidence_loc": "定位",
        "evidence_text": "原文",
    }


def _template_profile(template_version: str):
    if str(template_version or "").lower().startswith("v2"):
        return {
            "title_color": RGBColor(31, 78, 120),
            "heading_color": RGBColor(44, 62, 80),
            "table_header_fill": "CFE2F3",
            "accent_color": RGBColor(31, 78, 120),
        }
    return {
        "title_color": RGBColor(0, 0, 0),
        "heading_color": RGBColor(0, 0, 0),
        "table_header_fill": "D9E1F2",
        "accent_color": RGBColor(90, 90, 90),
    }


def _set_doc_layout(doc: Document):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)


def _apply_run_font(run, name: str, size: float, bold: bool = False, color: RGBColor | None = None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def _add_title(doc: Document, text: str, profile: dict):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    _apply_run_font(r, "SimHei", 16, bold=True, color=profile["title_color"])
    p.paragraph_format.space_after = Pt(12)


def _add_heading(doc: Document, text: str, level: int, profile: dict):
    p = doc.add_paragraph()
    r = p.add_run(text)
    if level == 2:
        _apply_run_font(r, "SimHei", 14, bold=True,
                        color=profile["heading_color"])
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(8)
    else:
        _apply_run_font(r, "SimHei", 12, bold=True,
                        color=profile["heading_color"])
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    return p


def _add_body(doc: Document, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    _apply_run_font(r, "SimSun", 12)
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(0.74)
    return p


def _set_cell_padding(cell, top=80, left=120, bottom=80, right=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_header_row_repeat(row):
    tr = row._tr
    tr_pr = tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tr_pr.append(tbl_header)


def _set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(
        rf'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{fill}"/>')
    tc_pr.append(shd)


def _add_toc_placeholder(doc: Document, labels: dict, profile: dict):
    _add_heading(doc, labels["toc"], level=2, profile=profile)
    p = doc.add_paragraph()
    r = p.add_run(labels["toc_hint"])
    _apply_run_font(r, "SimSun", 11, color=RGBColor(100, 100, 100))
    p.paragraph_format.space_after = Pt(8)
    p_field = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = labels["toc"]
    run.append(text)
    fld.append(run)
    p_field._p.append(fld)
    doc.add_page_break()


def _set_table_col_widths(table, width_cm_list):
    table.autofit = False
    for col_idx, w in enumerate(width_cm_list):
        for row in table.rows:
            row.cells[col_idx].width = Cm(w)


def _set_header_footer(doc: Document, contract_id: str, generated_at: str, template_version: str, labels: dict, brand: str):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = ""
    title = labels["header_title"] if not brand else f"{labels['header_title']} - {brand}"
    header_run = header.add_run(
        f"{title}  |  {labels['contract_id']}: {contract_id}")
    _apply_run_font(header_run, "SimSun", 9, color=RGBColor(90, 90, 90))
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT

    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer_run = footer.add_run(
        f"{labels['generated_at']}: {generated_at}  |  {labels['template_version']}: {template_version}")
    _apply_run_font(footer_run, "SimSun", 9, color=RGBColor(110, 110, 110))
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def render_tax_audit_docx(report: dict, output_path: str, template_version: str = "v1.0", locale: str = "zh-CN", brand: str = ""):
    doc = Document()
    _set_doc_layout(doc)
    labels = _labels(locale)
    profile = _template_profile(template_version)

    overview = report.get("overview") or {}
    risk_summary = report.get("risk_summary") or {}
    review_summary = report.get("review_summary") or {}
    risk_items = report.get("risk_items") or []
    evidence_items = report.get("evidence_items") or []
    contract_id = str(report.get("contract_id") or "")
    generated_at = str(report.get("generated_at") or "")
    _set_header_footer(doc, contract_id=contract_id,
                       generated_at=generated_at, template_version=template_version, labels=labels, brand=brand)

    _add_title(doc, labels["report_title"], profile=profile)
    _add_body(doc, f"{labels['contract_id']}：{contract_id}")
    _add_body(doc, f"{labels['generated_at']}：{generated_at}")
    _add_body(
        doc, f"{labels['contract_file']}：{overview.get('contract_filename', '')}")
    _add_toc_placeholder(doc, labels, profile=profile)
    _add_heading(doc, labels["overview"], level=2, profile=profile)
    _add_body(
        doc, f"{labels['parse_status']}：{overview.get('contract_parse_status', '')}")
    _add_body(
        doc, f"{labels['clause_count']}：{overview.get('clause_count', 0)}")
    _add_body(doc, f"{labels['issue_count']}：{overview.get('issue_count', 0)}")
    _add_body(doc, f"{labels['trace_count']}：{overview.get('trace_count', 0)}")

    _add_heading(doc, labels["summary"], level=2, profile=profile)
    _add_body(doc, f"{labels['high']}：{risk_summary.get('high', 0)}")
    _add_body(doc, f"{labels['medium']}：{risk_summary.get('medium', 0)}")
    _add_body(doc, f"{labels['low']}：{risk_summary.get('low', 0)}")
    _add_body(doc, f"{labels['pending']}：{review_summary.get('pending', 0)}")
    _add_body(
        doc, f"{labels['confirmed']}：{review_summary.get('confirmed', 0)}")

    _add_heading(doc, labels["risks"], level=2, profile=profile)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    _set_table_col_widths(table, [1.0, 1.8, 4.0, 5.2, 3.5])
    hdr = table.rows[0].cells
    _set_header_row_repeat(table.rows[0])
    titles = [
        labels["table_no"],
        labels["table_level"],
        labels["table_issue"],
        labels["table_evidence"],
        labels["table_suggestion"],
    ]
    for i, title in enumerate(titles):
        hdr[i].text = title
        p = hdr[i].paragraphs[0]
        for run in p.runs:
            _apply_run_font(run, "SimHei", 10.5, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_shading(hdr[i], profile["table_header_fill"])
        _set_cell_padding(hdr[i])
    evidence_map = {str(x.get("issue_id") or ""): x for x in evidence_items}
    for idx, item in enumerate(risk_items, start=1):
        row = table.add_row().cells
        issue_id = str(item.get("issue_id") or "")
        evidence = evidence_map.get(issue_id) or {}
        row[0].text = str(idx)
        row[1].text = str(item.get("risk_level") or "")
        row[2].text = str(item.get("issue_text") or "")
        row[3].text = str(evidence.get("source_text") or "")
        row[4].text = str(item.get("suggestion") or "")
        for col in range(5):
            p = row[col].paragraphs[0]
            for run in p.runs:
                _apply_run_font(run, "SimSun", 10.5)
            _set_cell_padding(row[col])

    _add_heading(doc, labels["appendix"], level=2, profile=profile)
    for idx, evidence in enumerate(evidence_items, start=1):
        _add_heading(
            doc, f"{labels['evidence_title']} {idx}", level=3, profile=profile)
        _add_body(
            doc, f"{labels['evidence_issue_id']}：{evidence.get('issue_id', '')}")
        _add_body(
            doc, f"{labels['evidence_law']}：{evidence.get('law_title', '')} {evidence.get('article_no', '')}")
        _add_body(
            doc, f"{labels['evidence_loc']}：第{evidence.get('source_page', '')}页 / 段落{evidence.get('source_paragraph', '')}")
        _add_body(
            doc, f"{labels['evidence_text']}：{evidence.get('source_text', '')}")

    _add_heading(doc, labels["risk_details"], level=2, profile=profile)
    evidence_map = {}
    for evidence in evidence_items:
        issue_key = str(evidence.get("issue_id") or "")
        evidence_map.setdefault(issue_key, []).append(evidence)
    for idx, item in enumerate(risk_items, start=1):
        issue_id = str(item.get("issue_id") or "")
        _add_heading(
            doc, f"{labels['risk_no']} R-{idx:03d}", level=3, profile=profile)
        _add_body(doc, f"{labels['risk_level']}：{item.get('risk_level', '')}")
        _add_body(doc, f"{labels['risk_text']}：{item.get('issue_text', '')}")
        _add_body(
            doc, f"{labels['risk_suggestion']}：{item.get('suggestion', '')}")
        for ev in evidence_map.get(issue_id, [])[:2]:
            p = doc.add_paragraph()
            r1 = p.add_run(
                f"{labels['evidence_loc']}：第{ev.get('source_page', '')}页 / 段落{ev.get('source_paragraph', '')} ")
            _apply_run_font(r1, "SimSun", 11, bold=True,
                            color=profile["accent_color"])
            r2 = p.add_run(str(ev.get("source_text") or ""))
            _apply_run_font(r2, "SimSun", 11)
            p.paragraph_format.space_after = Pt(6)

    doc.save(output_path)
