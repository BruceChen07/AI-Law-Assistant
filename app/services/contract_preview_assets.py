import hashlib
import json
import logging
import os
import re
import shutil
from datetime import datetime
from typing import Dict, Any, List, Tuple

from docx import Document as DocxDocument
from PIL import Image, ImageDraw, ImageFont
from pypdf import PdfReader

from app.core.mineru_ocr import run_mineru_extract
from app.core.utils import extract_text_with_config

try:
    from pdf2image import convert_from_path
except Exception:
    convert_from_path = None


logger = logging.getLogger("law_assistant")


def _looks_like_heading_line(text: str) -> bool:
    raw = str(text or "").strip()
    if not raw:
        return False
    if len(raw) <= 40 and re.match(r"^第[一二三四五六七八九十百千万0-9]+[章节条款]", raw):
        return True
    m = re.match(
        r"^([一二三四五六七八九十]+[、.．]|[0-9]+(?:\.[0-9]+){0,3}[、.．]?)(.*)$", raw)
    if m:
        tail = str(m.group(2) or "").strip()
        if not tail:
            return True
        if len(tail) <= 12 and not re.search(r"[：:，,。；;]", tail):
            return True
        return False
    if len(raw) <= 28 and re.match(r"^[（(]?[一二三四五六七八九十0-9]+[)）][^。；;，,:：]{0,24}$", raw):
        return True
    if len(raw) <= 24 and not re.search(r"[。；;，,:：]", raw) and re.search(r"(合同|条款|价款|支付|期限|服务|违约|保密|发票|税率)", raw):
        return True
    return False


def _safe_int(v: Any, default: int) -> int:
    try:
        return int(v)
    except Exception:
        return default


def _safe_float(v: Any, default: float) -> float:
    try:
        return float(v)
    except Exception:
        return float(default)


def _resolve_cache_dir(cfg: Dict[str, Any], preview_cfg: Dict[str, Any]) -> str:
    raw = str(preview_cfg.get("cache_dir") or "").strip()
    if not raw:
        return os.path.join(str(cfg.get("files_dir") or ""), "preview_cache")
    if os.path.isabs(raw):
        return raw
    return os.path.abspath(os.path.join(str(cfg.get("files_dir") or ""), raw))


def _file_signature(path: str) -> str:
    st = os.stat(path)
    size = int(st.st_size or 0)
    sample_size = 1024 * 1024
    head_size = min(size, sample_size)
    tail_size = min(size, sample_size)
    h = hashlib.sha1()
    h.update(str(size).encode("utf-8"))
    with open(path, "rb") as f:
        if head_size > 0:
            h.update(f.read(head_size))
        if size > tail_size:
            f.seek(max(0, size - tail_size))
            h.update(f.read(tail_size))
    return h.hexdigest()


def _settings_signature(settings: Dict[str, Any]) -> str:
    payload = json.dumps(settings or {}, sort_keys=True,
                         ensure_ascii=False, separators=(",", ":"))
    return hashlib.sha1(payload.encode("utf-8")).hexdigest()


def _build_cache_key(file_signature: str, settings_signature: str, ext: str) -> str:
    seed = f"{file_signature}|{settings_signature}|{ext}"
    return hashlib.sha1(seed.encode("utf-8")).hexdigest()


def _is_valid_preview_image(path: str) -> bool:
    if not path or not os.path.exists(path):
        return False
    try:
        if os.path.getsize(path) < 64:
            return False
        with open(path, "rb") as f:
            sig = f.read(8)
        return sig == b"\x89PNG\r\n\x1a\n"
    except Exception:
        return False


def _build_text_pages(text: str, lines_per_page: int) -> List[Dict[str, Any]]:
    lines = [ln for ln in str(text or "").splitlines()]
    if not lines:
        return [{"page_no": 1, "width": 0, "height": 0, "blocks": [], "image_file": ""}]
    size = max(1, _safe_int(lines_per_page, 80))
    pages: List[Dict[str, Any]] = []
    for idx in range(0, len(lines), size):
        chunk = lines[idx: idx + size]
        blocks = []
        denom = max(len(chunk), 1)
        for j, line in enumerate(chunk, 1):
            txt = str(line).strip()
            if not txt:
                continue
            y = (j - 1) / denom
            h = 1 / denom
            blocks.append({
                "block_id": f"p{(idx // size) + 1}-l{j}",
                "text": txt,
                "bbox": [0.04, round(y, 6), 0.92, round(h, 6)],
                "is_heading": _looks_like_heading_line(txt),
            })
        pages.append({
            "page_no": (idx // size) + 1,
            "width": 0,
            "height": 0,
            "blocks": blocks,
            "image_file": "",
        })
    return pages


def _render_pdf_pages(file_path: str, out_dir: str, dpi: int, max_pages: int) -> List[Dict[str, Any]]:
    if convert_from_path is None:
        raise RuntimeError("pdf2image not available")
    images = convert_from_path(
        file_path,
        dpi=max(72, _safe_int(dpi, 160)),
        first_page=1,
        last_page=max(1, _safe_int(max_pages, 30)),
        fmt="png",
    )
    pages: List[Dict[str, Any]] = []
    for i, image in enumerate(images, 1):
        image_file = os.path.join(out_dir, f"page_{i}.png")
        image.save(image_file, "PNG")
        pages.append({
            "page_no": i,
            "width": int(getattr(image, "width", 0) or 0),
            "height": int(getattr(image, "height", 0) or 0),
            "blocks": [],
            "image_file": image_file,
        })
    return pages


def _extract_pdf_text_blocks(file_path: str, max_pages: int) -> Tuple[Dict[int, List[Dict[str, Any]]], int]:
    reader = PdfReader(file_path)
    total = min(len(reader.pages), max(1, _safe_int(max_pages, 30)))
    out: Dict[int, List[Dict[str, Any]]] = {}
    for i in range(total):
        text = reader.pages[i].extract_text() or ""
        lines = [ln.strip() for ln in text.splitlines() if ln and ln.strip()]
        blocks = []
        denom = max(len(lines), 1)
        for j, line in enumerate(lines, 1):
            y = (j - 1) / denom
            h = 1 / denom
            blocks.append({
                "block_id": f"p{i + 1}-l{j}",
                "text": line,
                "bbox": [0.04, round(y, 6), 0.92, round(h, 6)],
                "is_heading": _looks_like_heading_line(line),
            })
        out[i + 1] = blocks
    return out, total


def _load_json(path: str) -> Any:
    if not path or not os.path.exists(path):
        return None
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def _normalize_bbox(raw_bbox: Any) -> Tuple[float, float, float, float]:
    if not isinstance(raw_bbox, list) or len(raw_bbox) < 4:
        return 0.0, 0.0, 0.0, 0.0
    x1 = _safe_float(raw_bbox[0], 0.0)
    y1 = _safe_float(raw_bbox[1], 0.0)
    x2 = _safe_float(raw_bbox[2], x1)
    y2 = _safe_float(raw_bbox[3], y1)
    left = min(x1, x2)
    top = min(y1, y2)
    right = max(x1, x2)
    bottom = max(y1, y2)
    return left, top, right, bottom


def _normalize_ratio_bbox(x: float, y: float, w: float, h: float, page_w: float, page_h: float) -> List[float]:
    safe_w = max(1.0, page_w)
    safe_h = max(1.0, page_h)
    rx = min(max(x / safe_w, 0.0), 1.0)
    ry = min(max(y / safe_h, 0.0), 1.0)
    rw = min(max(w / safe_w, 0.0), 1.0)
    rh = min(max(h / safe_h, 0.0), 1.0)
    return [round(rx, 6), round(ry, 6), round(rw, 6), round(rh, 6)]


def _build_blocks_from_mineru(file_path: str, out_dir: str, max_pages: int) -> Tuple[Dict[int, List[Dict[str, Any]]], Dict[int, Dict[str, Any]], Dict[str, Any]]:
    mineru_dir = os.path.join(out_dir, "mineru_layout")
    result = run_mineru_extract(file_path, output_dir=mineru_dir)
    middle_path = str(result.get("middle_path") or "")
    content_list_path = str(result.get("content_list_path") or "")
    middle = _load_json(middle_path)
    content_list = _load_json(content_list_path)
    pdf_info = middle.get("pdf_info") if isinstance(
        middle, dict) and isinstance(middle.get("pdf_info"), list) else []
    pages_meta: Dict[int, Dict[str, Any]] = {}
    for idx, page in enumerate(pdf_info, 1):
        if idx > max_pages:
            break
        page_size = page.get("page_size") if isinstance(page, dict) else None
        if isinstance(page_size, list) and len(page_size) >= 2:
            w = max(1.0, _safe_float(page_size[0], 595.0))
            h = max(1.0, _safe_float(page_size[1], 842.0))
        else:
            w, h = 595.0, 842.0
        pages_meta[idx] = {"page_width_pt": w, "page_height_pt": h}
    rows = content_list if isinstance(content_list, list) else []
    page_items: Dict[int, List[Dict[str, Any]]] = {}
    for idx, row in enumerate(rows, 1):
        if not isinstance(row, dict):
            continue
        page_idx = _safe_int(row.get("page_idx"), -1)
        page_no = page_idx + 1
        if page_no <= 0 or page_no > max_pages:
            continue
        item_type = str(row.get("type") or "")
        if item_type in {"discarded"}:
            continue
        text = str(row.get("text") or "").strip()
        if item_type == "text" and not text:
            continue
        bbox = row.get("bbox")
        if not isinstance(bbox, list) or len(bbox) < 4:
            continue
        page_items.setdefault(page_no, []).append({"row_idx": idx, "row": row})
    blocks_map: Dict[int, List[Dict[str, Any]]] = {}
    for page_no, items in page_items.items():
        meta = pages_meta.get(page_no) or {
            "page_width_pt": 595.0, "page_height_pt": 842.0}
        page_w = _safe_float(meta.get("page_width_pt"), 595.0)
        page_h = _safe_float(meta.get("page_height_pt"), 842.0)
        max_x2 = 0.0
        max_y2 = 0.0
        for item in items:
            x1, y1, x2, y2 = _normalize_bbox(item["row"].get("bbox"))
            max_x2 = max(max_x2, x2)
            max_y2 = max(max_y2, y2)
        scale_x = 1.0
        scale_y = 1.0
        if page_w > 0 and max_x2 > page_w * 1.2:
            scale_x = max_x2 / page_w
        if page_h > 0 and max_y2 > page_h * 1.2:
            scale_y = max_y2 / page_h
        if abs(scale_x - scale_y) <= 0.3:
            scale_x = max(scale_x, 1.0)
            scale_y = scale_x
        blocks: List[Dict[str, Any]] = []
        for item in items:
            row = item["row"]
            x1, y1, x2, y2 = _normalize_bbox(row.get("bbox"))
            left_pt = x1 / max(1.0, scale_x)
            top_pt = y1 / max(1.0, scale_y)
            width_pt = max(0.0, (x2 - x1) / max(1.0, scale_x))
            height_pt = max(0.0, (y2 - y1) / max(1.0, scale_y))
            text = str(row.get("text") or "").strip()
            text_level = _safe_int(row.get("text_level"), 9)
            block_type = str(row.get("type") or "text")
            is_heading = bool(text_level in {1, 2}) if text else False
            if text and not is_heading:
                is_heading = _looks_like_heading_line(text)
            block = {
                "block_id": f"p{page_no}-m{_safe_int(item.get('row_idx'), len(blocks) + 1)}",
                "text": text,
                "block_type": block_type,
                "bbox": _normalize_ratio_bbox(left_pt, top_pt, width_pt, height_pt, page_w, page_h),
                "bbox_pt": [round(left_pt, 3), round(top_pt, 3), round(width_pt, 3), round(height_pt, 3)],
                "coord_unit": "pt",
                "coord_origin": "top_left",
                "is_heading": is_heading,
                "confidence": _safe_float(row.get("score"), 1.0),
            }
            blocks.append(block)
        blocks_map[page_no] = blocks
    block_total = sum(len(v) for v in blocks_map.values())
    markdown_text_length = len(str(result.get("text") or ""))
    return blocks_map, pages_meta, {
        "md_path": str(result.get("md_path") or ""),
        "middle_path": middle_path,
        "content_list_path": content_list_path,
        "page_count": _safe_int(result.get("pages"), len(blocks_map)),
        "block_count": block_total,
        "markdown_text_length": markdown_text_length,
        "coord_unit": "pt",
        "coord_origin": "top_left",
    }


def _extract_docx_plain_text(file_path: str) -> str:
    try:
        doc = DocxDocument(file_path)
    except Exception:
        return ""
    lines = []
    for p in doc.paragraphs:
        txt = _normalize_line_text(getattr(p, "text", ""))
        if txt:
            lines.append(txt)
    return "\n".join(lines)


def _convert_docx_to_pdf_with_docx2pdf(docx_path: str, output_pdf: str) -> str:
    from docx2pdf import convert

    convert(docx_path, output_pdf)
    if not os.path.exists(output_pdf):
        raise RuntimeError("docx2pdf conversion did not produce output")
    return output_pdf


def _convert_docx_to_pdf_with_win32com(docx_path: str, output_pdf: str) -> str:
    import win32com.client

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    try:
        document = word.Documents.Open(os.path.abspath(docx_path))
        document.SaveAs(os.path.abspath(output_pdf), FileFormat=17)
        document.Close(False)
    finally:
        word.Quit()
    if not os.path.exists(output_pdf):
        raise RuntimeError("win32com conversion did not produce output")
    return output_pdf


def _convert_docx_to_pdf(docx_path: str, output_pdf: str) -> Tuple[str, str]:
    os.makedirs(os.path.dirname(output_pdf), exist_ok=True)
    if os.path.exists(output_pdf):
        try:
            os.remove(output_pdf)
        except Exception:
            pass
    tried: List[str] = []
    for method in ("docx2pdf", "win32com"):
        try:
            if method == "docx2pdf":
                return _convert_docx_to_pdf_with_docx2pdf(docx_path, output_pdf), method
            return _convert_docx_to_pdf_with_win32com(docx_path, output_pdf), method
        except Exception:
            tried.append(method)
    raise RuntimeError(f"docx to pdf conversion failed via {','.join(tried)}")


def _docx_pdf_quality_gate(
    docx_text: str,
    mineru_meta: Dict[str, Any],
    min_text_ratio: float,
    min_docx_chars: int,
) -> Tuple[bool, Dict[str, Any]]:
    docx_chars = len(str(docx_text or ""))
    mineru_chars = _safe_int(mineru_meta.get("markdown_text_length"), 0)
    page_count = _safe_int(mineru_meta.get("page_count"), 0)
    block_count = _safe_int(mineru_meta.get("block_count"), 0)
    ratio = 1.0 if docx_chars <= 0 else (mineru_chars / max(docx_chars, 1))
    enough_text = True
    if docx_chars >= max(20, _safe_int(min_docx_chars, 120)):
        enough_text = ratio >= max(0.05, _safe_float(min_text_ratio, 0.35))
    passed = page_count > 0 and block_count > 0 and enough_text
    return passed, {
        "docx_chars": docx_chars,
        "mineru_chars": mineru_chars,
        "text_ratio": round(ratio, 4),
        "page_count": page_count,
        "block_count": block_count,
    }


def _load_preview_font(size: int) -> Any:
    candidates = [
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simsun.ttc",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    ]
    for path in candidates:
        try:
            if os.path.exists(path):
                return ImageFont.truetype(path, size=size)
        except Exception:
            pass
    return ImageFont.load_default()


def _normalize_line_text(text: str) -> str:
    return " ".join(str(text or "").replace("\t", " ").split()).strip()


def _wrap_text_by_width(text: str, draw: Any, font: Any, max_width: int, hard_limit: int) -> List[str]:
    raw = _normalize_line_text(text)
    if not raw:
        return []
    limit = max(10, _safe_int(hard_limit, 44))
    out: List[str] = []
    buf = ""
    for ch in raw:
        nxt = f"{buf}{ch}"
        if len(nxt) <= limit and float(draw.textlength(nxt, font=font)) <= max_width:
            buf = nxt
            continue
        if buf.strip():
            out.append(buf.rstrip())
        buf = "" if ch == " " else ch
    if buf.strip():
        out.append(buf.rstrip())
    return out


def _paragraph_profile(paragraph: Any) -> Dict[str, Any]:
    style_name = str(getattr(getattr(paragraph, "style", None),
                     "name", "") or "").strip().lower()
    txt = _normalize_line_text(getattr(paragraph, "text", ""))
    is_heading = style_name.startswith("heading") or style_name.startswith(
        "标题") or _looks_like_heading_line(txt)
    level = 1
    m = re.search(r"(\d+)", style_name)
    if m:
        level = max(1, min(6, _safe_int(m.group(1), 1)))
    if is_heading:
        scale = max(1.0, 1.34 - (level - 1) * 0.07)
        return {"line_scale": scale, "gap_before": 12 if level <= 2 else 8, "gap_after": 8 if level <= 2 else 6, "indent": 0, "is_heading": True}
    is_list = style_name.startswith("list") or txt.startswith(
        ("-", "•", "*", "1.", "2.", "3."))
    if is_list:
        return {"line_scale": 1.0, "gap_before": 2, "gap_after": 4, "indent": 24, "is_heading": False}
    return {"line_scale": 1.0, "gap_before": 2, "gap_after": 8, "indent": 0, "is_heading": False}


def _render_docx_pages(file_path: str, out_dir: str, max_pages: int, page_width: int, page_height: int, margin: int, line_height: int, chars_per_line: int, paragraph_spacing: int) -> Tuple[List[Dict[str, Any]], int]:
    doc = DocxDocument(file_path)
    line_h = max(16, _safe_int(line_height, 36))
    base_gap = max(4, _safe_int(paragraph_spacing, int(line_h * 0.45)))
    raw_margin = _safe_int(margin, 72)
    safe_margin = min(max(24, raw_margin), max(
        36, page_width // 6), max(36, page_height // 8))
    font_cache: Dict[int, Any] = {}

    def _font_for(scale: float) -> Any:
        size = max(12, int(line_h * 0.72 * max(0.85, float(scale))))
        if size not in font_cache:
            font_cache[size] = _load_preview_font(size)
        return font_cache[size]

    measure_img = Image.new("RGB", (8, 8), color=(255, 255, 255))
    measure = ImageDraw.Draw(measure_img)
    flow: List[Dict[str, Any]] = []
    for p in doc.paragraphs:
        txt = _normalize_line_text(p.text)
        profile = _paragraph_profile(p)
        if not txt:
            if flow and flow[-1].get("kind") != "gap":
                flow.append({"kind": "gap", "h": base_gap})
            continue
        line_scale = float(profile.get("line_scale") or 1.0)
        indent = max(0, _safe_int(profile.get("indent"), 0))
        line_h_local = max(16, int(line_h * max(0.9, line_scale)))
        gap_before = max(0, _safe_int(profile.get("gap_before"), 0))
        gap_after = max(0, _safe_int(profile.get("gap_after"), base_gap))
        font = _font_for(line_scale)
        max_text_width = max(120, page_width - safe_margin * 2 - indent)
        hard_limit = max(
            10, int(chars_per_line * (1.08 if line_scale > 1.08 else 1.0)))
        wrapped = _wrap_text_by_width(
            txt, measure, font, max_text_width, hard_limit)
        if flow and gap_before > 0:
            flow.append({"kind": "gap", "h": gap_before})
        for ln in wrapped:
            flow.append({"kind": "line", "text": ln, "line_h": line_h_local, "line_scale": line_scale,
                        "indent": indent, "max_w": max_text_width, "is_heading": bool(profile.get("is_heading"))})
        flow.append({"kind": "gap", "h": gap_after})
    while flow and flow[-1].get("kind") == "gap":
        flow.pop()

    line_total = sum(1 for it in flow if it.get("kind") == "line")
    if line_total == 0:
        return [{"page_no": 1, "width": page_width, "height": page_height, "blocks": [], "image_file": ""}], 0

    pages: List[Dict[str, Any]] = []
    cursor = 0
    page_no = 1
    max_page_count = max(1, _safe_int(max_pages, 30))
    while cursor < len(flow) and page_no <= max_page_count:
        image = Image.new("RGB", (page_width, page_height),
                          color=(255, 255, 255))
        draw = ImageDraw.Draw(image)
        y = safe_margin
        line_idx = 0
        blocks = []
        while cursor < len(flow):
            item = flow[cursor]
            kind = str(item.get("kind") or "")
            if kind == "gap":
                h = max(0, _safe_int(item.get("h"), base_gap))
                if blocks and y + h <= page_height - safe_margin:
                    y += h
                cursor += 1
                continue
            local_h = max(16, _safe_int(item.get("line_h"), line_h))
            if y + local_h > page_height - safe_margin:
                break
            line = str(item.get("text") or "")
            line_scale = float(item.get("line_scale") or 1.0)
            indent = max(0, _safe_int(item.get("indent"), 0))
            max_w = max(120, _safe_int(item.get("max_w"),
                        page_width - safe_margin * 2 - indent))
            x = safe_margin + indent
            draw.text((x, y), line, fill=(25, 25, 25),
                      font=_font_for(line_scale))
            line_idx += 1
            blocks.append({
                "block_id": f"p{page_no}-l{line_idx}",
                "text": line,
                "bbox": [round(x / page_width, 6), round(y / page_height, 6), round(max_w / page_width, 6), round(local_h / page_height, 6)],
                "is_heading": bool(item.get("is_heading")),
            })
            y += local_h
            cursor += 1
        image_file = os.path.join(out_dir, f"page_{page_no}.png")
        image.save(image_file, "PNG")
        pages.append({"page_no": page_no, "width": page_width,
                     "height": page_height, "blocks": blocks, "image_file": image_file})
        page_no += 1
    return pages, line_total


def build_contract_preview_manifest(cfg: Dict[str, Any], document_id: str, file_path: str, mime_type: str = "") -> Dict[str, Any]:
    if not file_path or not os.path.exists(file_path):
        raise FileNotFoundError("document file not found")

    preview_cfg = cfg.get("contract_preview") if isinstance(
        cfg.get("contract_preview"), dict) else {}
    enabled = bool(preview_cfg.get("enabled", True))
    dpi = _safe_int(preview_cfg.get("pdf_dpi"), 160)
    max_pages = _safe_int(preview_cfg.get("pdf_max_pages"), 30)
    lines_per_page = _safe_int(preview_cfg.get("text_lines_per_page"), 80)
    docx_visual_enabled = bool(preview_cfg.get("docx_visual_enabled", True))
    docx_page_width = _safe_int(preview_cfg.get("docx_page_width"), 1240)
    docx_page_height = _safe_int(preview_cfg.get("docx_page_height"), 1754)
    docx_margin = _safe_int(preview_cfg.get("docx_margin"), 72)
    docx_line_height = _safe_int(preview_cfg.get("docx_line_height"), 36)
    docx_chars_per_line = _safe_int(preview_cfg.get("docx_chars_per_line"), 44)
    docx_paragraph_spacing = _safe_int(
        preview_cfg.get("docx_paragraph_spacing"), 14)
    coord_provider = str(preview_cfg.get("coord_provider")
                         or "mineru").strip().lower()
    docx_coord_from_pdf = bool(preview_cfg.get("docx_coord_from_pdf", True))
    docx_pdf_min_text_ratio = _safe_float(
        preview_cfg.get("docx_pdf_min_text_ratio"), 0.35)
    docx_pdf_gate_min_chars = _safe_int(
        preview_cfg.get("docx_pdf_gate_min_chars"), 120)
    docx_pdf_keep_file = bool(preview_cfg.get("docx_pdf_keep_file", False))
    ext = os.path.splitext(file_path)[1].lower()

    cache_root = _resolve_cache_dir(cfg, preview_cfg)
    os.makedirs(cache_root, exist_ok=True)

    signature = _file_signature(file_path)
    settings = {
        "preview_schema_version": 4,
        "pdf_dpi": dpi,
        "pdf_max_pages": max_pages,
        "text_lines_per_page": lines_per_page,
        "docx_visual_enabled": docx_visual_enabled,
        "docx_page_width": docx_page_width,
        "docx_page_height": docx_page_height,
        "docx_margin": docx_margin,
        "docx_line_height": docx_line_height,
        "docx_chars_per_line": docx_chars_per_line,
        "docx_paragraph_spacing": docx_paragraph_spacing,
        "coord_provider": coord_provider,
        "docx_coord_from_pdf": docx_coord_from_pdf,
        "docx_pdf_min_text_ratio": docx_pdf_min_text_ratio,
        "docx_pdf_gate_min_chars": docx_pdf_gate_min_chars,
        "docx_pdf_keep_file": docx_pdf_keep_file,
    }
    settings_sig = _settings_signature(settings)
    cache_key = _build_cache_key(signature, settings_sig, ext)
    cache_dir = os.path.join(cache_root, cache_key)
    manifest_path = os.path.join(cache_dir, "manifest.json")

    if os.path.exists(manifest_path):
        try:
            with open(manifest_path, "r", encoding="utf-8") as f:
                cached = json.load(f)
            same = (
                str(cached.get("file_signature") or "") == signature
                and cached.get("settings") == settings
            )
            if same:
                valid = True
                for p in cached.get("pages") or []:
                    image_file = str(p.get("image_file") or "")
                    if image_file and not _is_valid_preview_image(image_file):
                        valid = False
                        break
                if valid:
                    cached["document_id"] = document_id
                    cached["mime_type"] = str(
                        mime_type or cached.get("mime_type") or "")
                    return cached
        except Exception:
            pass

    if os.path.exists(cache_dir):
        shutil.rmtree(cache_dir, ignore_errors=True)
    os.makedirs(cache_dir, exist_ok=True)

    mode = "text"
    source = "text_fallback"
    pages: List[Dict[str, Any]] = []
    text = ""
    meta = {
        "ocr_used": False,
        "ocr_engine": "",
        "line_total": 0,
        "page_count": 0,
        "coord_provider": "ratio_estimate",
        "coord_unit": "ratio",
        "coord_origin": "top_left",
        "coord_source": "",
        "docx_pdf_conversion": {},
    }

    if enabled and ext == ".pdf":
        try:
            pages = _render_pdf_pages(file_path, cache_dir, dpi, max_pages)
            blocks_map: Dict[int, List[Dict[str, Any]]] = {}
            page_count = len(pages)
            if coord_provider == "mineru":
                try:
                    blocks_map, pages_meta, mineru_meta = _build_blocks_from_mineru(
                        file_path=file_path,
                        out_dir=cache_dir,
                        max_pages=max_pages,
                    )
                    if blocks_map:
                        page_count = _safe_int(
                            mineru_meta.get("page_count"), page_count)
                        meta["coord_provider"] = "mineru"
                        meta["coord_unit"] = "pt"
                        meta["coord_origin"] = "top_left"
                        meta["coord_source"] = "mineru_content_list"
                        for p in pages:
                            page_no = _safe_int(p.get("page_no"), 0)
                            if page_no > 0:
                                page_meta = pages_meta.get(page_no) or {}
                                p["coord_unit"] = "pt"
                                p["coord_origin"] = "top_left"
                                p["page_width_pt"] = round(_safe_float(
                                    page_meta.get("page_width_pt"), 0.0), 3)
                                p["page_height_pt"] = round(_safe_float(
                                    page_meta.get("page_height_pt"), 0.0), 3)
                except Exception:
                    logger.exception(
                        "preview_mineru_layout_build_failed file=%s", file_path)
            if not blocks_map:
                blocks_map, page_count = _extract_pdf_text_blocks(
                    file_path, max_pages)
            for p in pages:
                page_no = int(p.get("page_no") or 0)
                p["blocks"] = blocks_map.get(page_no, [])
                if not p.get("coord_unit"):
                    p["coord_unit"] = "ratio"
                    p["coord_origin"] = "top_left"
            mode = "visual"
            source = "pdf_raster_mineru" if meta.get(
                "coord_provider") == "mineru" else "pdf_raster"
            meta["page_count"] = page_count
            meta["line_total"] = sum(len(p.get("blocks") or []) for p in pages)
        except Exception:
            logger.exception("preview_visual_build_failed file=%s", file_path)

    if enabled and mode != "visual" and ext == ".docx" and docx_visual_enabled:
        if coord_provider == "mineru" and docx_coord_from_pdf:
            try:
                converted_pdf = os.path.join(
                    cache_dir, "docx_coord_source.pdf")
                converted_pdf, conversion_method = _convert_docx_to_pdf(
                    docx_path=file_path,
                    output_pdf=converted_pdf,
                )
                pages = _render_pdf_pages(
                    converted_pdf, cache_dir, dpi, max_pages)
                blocks_map, pages_meta, mineru_meta = _build_blocks_from_mineru(
                    file_path=converted_pdf,
                    out_dir=cache_dir,
                    max_pages=max_pages,
                )
                gate_ok, gate_info = _docx_pdf_quality_gate(
                    docx_text=_extract_docx_plain_text(file_path),
                    mineru_meta=mineru_meta,
                    min_text_ratio=docx_pdf_min_text_ratio,
                    min_docx_chars=docx_pdf_gate_min_chars,
                )
                if gate_ok and blocks_map:
                    for p in pages:
                        page_no = _safe_int(p.get("page_no"), 0)
                        p["blocks"] = blocks_map.get(page_no, [])
                        page_meta = pages_meta.get(page_no) or {}
                        p["coord_unit"] = "pt"
                        p["coord_origin"] = "top_left"
                        p["page_width_pt"] = round(_safe_float(
                            page_meta.get("page_width_pt"), 0.0), 3)
                        p["page_height_pt"] = round(_safe_float(
                            page_meta.get("page_height_pt"), 0.0), 3)
                    mode = "visual"
                    source = "docx_pdf_mineru"
                    meta["page_count"] = _safe_int(
                        mineru_meta.get("page_count"), len(pages))
                    meta["line_total"] = sum(
                        len(p.get("blocks") or []) for p in pages)
                    meta["coord_provider"] = "mineru"
                    meta["coord_unit"] = "pt"
                    meta["coord_origin"] = "top_left"
                    meta["coord_source"] = "docx_to_pdf_mineru_content_list"
                    meta["docx_pdf_conversion"] = {
                        "method": conversion_method,
                        "quality": gate_info,
                    }
                else:
                    meta["docx_pdf_conversion"] = {
                        "method": conversion_method,
                        "quality": gate_info,
                        "fallback": "docx_raster",
                    }
                if not docx_pdf_keep_file and os.path.exists(converted_pdf):
                    try:
                        os.remove(converted_pdf)
                    except Exception:
                        pass
            except Exception:
                logger.exception(
                    "preview_docx_mineru_build_failed file=%s", file_path)
        try:
            if mode != "visual":
                pages, line_total = _render_docx_pages(
                    file_path=file_path,
                    out_dir=cache_dir,
                    max_pages=max_pages,
                    page_width=docx_page_width,
                    page_height=docx_page_height,
                    margin=docx_margin,
                    line_height=docx_line_height,
                    chars_per_line=docx_chars_per_line,
                    paragraph_spacing=docx_paragraph_spacing,
                )
                has_image = any(str(p.get("image_file") or "").strip()
                                for p in pages)
                if has_image:
                    mode = "visual"
                    source = "docx_raster"
                    meta["page_count"] = len(pages)
                    meta["line_total"] = int(line_total)
                    meta["coord_provider"] = "docx_layout"
                    meta["coord_unit"] = "ratio"
                    meta["coord_origin"] = "top_left"
                    meta["coord_source"] = "docx_visual_renderer"
        except Exception:
            logger.exception(
                "preview_docx_visual_build_failed file=%s", file_path)

    if mode != "visual":
        text, text_meta = extract_text_with_config(cfg, file_path)
        pages = _build_text_pages(text, lines_per_page)
        meta["ocr_used"] = bool(text_meta.get("ocr_used"))
        meta["ocr_engine"] = str(text_meta.get("ocr_engine") or "")
        meta["line_total"] = len(str(text).splitlines())
        meta["page_count"] = int(text_meta.get("page_count") or len(pages))
        meta["coord_provider"] = "text_fallback"
        meta["coord_unit"] = "ratio"
        meta["coord_origin"] = "top_left"
        meta["coord_source"] = "line_splitter"

    manifest = {
        "document_id": document_id,
        "mode": mode,
        "source": source,
        "mime_type": str(mime_type or ""),
        "ext": ext,
        "generated_at": datetime.utcnow().isoformat(),
        "file_signature": signature,
        "settings": settings,
        "cache_key": cache_key,
        "meta": {
            "page_total": len(pages),
            "page_count": int(meta.get("page_count") or len(pages)),
            "line_total": int(meta.get("line_total") or 0),
            "ocr_used": bool(meta.get("ocr_used")),
            "ocr_engine": str(meta.get("ocr_engine") or ""),
            "coord_provider": str(meta.get("coord_provider") or ""),
            "coord_unit": str(meta.get("coord_unit") or ""),
            "coord_origin": str(meta.get("coord_origin") or ""),
            "coord_source": str(meta.get("coord_source") or ""),
            "docx_pdf_conversion": meta.get("docx_pdf_conversion") if isinstance(meta.get("docx_pdf_conversion"), dict) else {},
        },
        "pages": pages,
        "text": text if mode == "text" else "",
    }

    with open(manifest_path, "w", encoding="utf-8") as f:
        json.dump(manifest, f, ensure_ascii=False, indent=2)
    return manifest


def find_preview_page(manifest: Dict[str, Any], page_no: int) -> Dict[str, Any]:
    target = int(page_no)
    for p in manifest.get("pages") or []:
        if int(p.get("page_no") or 0) == target:
            return p
    raise ValueError("preview page not found")
